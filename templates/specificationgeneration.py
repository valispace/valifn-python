from typing import Any, Dict
import operator
import warnings
import os
import getpass
import requests

from ast import Str
from collections.abc import Callable
from xmlrpc.client import Boolean
from valispace import API
from PIL import Image
from docx import Document
from docx.shared import Inches
from mailmerge import MailMerge
from natsort import natsorted 
from io import BytesIO


import urllib.request

from htmldocx import HtmlToDocx

#TEMPLATE_FILES_DIRECROTY = os.environ['USERPROFILE']+'/Documents/Generic Specification Template.docx' 

CURRENT_SPECIFICATION = {
    "id": 0,
    "name": "",
    "section": 0,
}

VALISPACE = {
    'domain': 'https://demonstration.valispace.com/',
    'username': 'AutomationsAPI',
    'password': 'AutomationsAPI'
}

DEFAULT_VALUES = {
    "project": 24    
}

def keep_tables_on_one_page(doc):
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
        ppr = tag.get_or_add_pPr()
        ppr.keepNext_val = True

def remove_all_empty_headings(document):
    for paragraph in document.paragraphs:
        if (paragraph.style.name == "Heading 1" or paragraph.style.name == "Heading 2") and paragraph.text =='':
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

def remove_all_but_last_section(document):
    sectPrs = document._element.xpath(".//w:pPr/w:sectPr")
    for sectPr in sectPrs:
        sectPr.getparent().remove(sectPr)
    return

def get_requirements_without_section(all_spec_requirements):
    unsorted_section_requirements = {}    
    no_section_requirements = {}
    for requirement in all_spec_requirements:
        if all_spec_requirements[requirement]['group'] == None:
            unsorted_section_requirements[requirement] =  all_spec_requirements[requirement]   

    sorted_requirements = natsorted(unsorted_section_requirements.items(), key=lambda x:(operator.getitem(x[1],'identifier').replace(". "," ")))
    
    for sorted_ in sorted_requirements:
        no_section_requirements[sorted_[0]] = sorted_[1] 

    return no_section_requirements
 
def get_section_requirements(all_specification_requirements,section):
    unsorted_section_requirements = {}
    section_requirements = {}
    for requirement in all_specification_requirements:
        if all_specification_requirements[requirement]['group'] == section:
            unsorted_section_requirements[requirement] =  all_specification_requirements[requirement]  
    
    sorted_requirements_by_section = natsorted(unsorted_section_requirements.items(), key=lambda x:(operator.getitem(x[1],'identifier').replace(". "," ")))
    
    for sorted_ in sorted_requirements_by_section:
        section_requirements[sorted_[0]] = sorted_[1] 

    return section_requirements

def get_specification_sections(api, project):
    specification_sections = {}
    all_specifications_groups = get_map(api, f"requirements/groups/?project={project}", "id", None, filter_specification)
    pre_sorted_section_groups = natsorted(all_specifications_groups.items(), key=lambda x:(operator.getitem(x[1],'name').replace(". "," ")))
    
    for pre_sorted in pre_sorted_section_groups:
        specification_sections[pre_sorted[0]] = pre_sorted[1] 

    return specification_sections

def put_images(document, all_project_images):

    for paragraph in document.paragraphs:
        if "Images_Placeholder" in paragraph.text:
            #still needs testing/refined. Currently only working/tested for tables!
            req_images = get_requirement_images(requirement_id,all_project_images)
            for image in req_images:
                image_data = Image.open(requests.get(req_images[image]['download_url'], stream=True).raw)    
                run = paragraph.add_run()
                run.add_picture(image_data.fp, height=Inches(2.0))    
                for run in reversed(list(paragraph.runs)[:-len(req_images)]):
                    paragraph._p.remove(run._r)
        elif "No_Images" in paragraph.text:
            p_el = paragraph._element
            p_el.getparent().remove(p_el)
            paragraph._p = paragraph._element = None

    for tables in document.tables:
            for celda in tables._cells:
                hasImages = False
                requirement_id = 0
                for p in celda.paragraphs:
                    if "Images_Placeholder" in p.text:
                       requirement_id = int(p.text[19:])
                       hasImages = True                       
                       p_el = p._element
                       p_el.getparent().remove(p_el)
                       p._p = p._element = None
                    elif "No_Images" in p.text:
                        p_el = p._element
                        p_el.getparent().remove(p_el)
                        p._p = p._element = None
                if hasImages:
                    new_image_p = celda.add_paragraph()
                    req_images = get_requirement_images(requirement_id,all_project_images)
                    for image in req_images:
                        image_data = Image.open(requests.get(req_images[image]['download_url'], stream=True).raw)    
                        run = new_image_p.add_run()
                        run.add_picture(image_data.fp, height=Inches(2.0))    
                    for run in reversed(list(new_image_p.runs)[:-len(req_images)]):
                        new_image_p._p.remove(run._r)
 
def clone_run_props(tmpl_run, this_run):
    this_run.bold = tmpl_run.bold
    this_run.italic = tmpl_run.italic
    this_run.underline = tmpl_run.underline
    this_run.font.color.rgb = tmpl_run.font.color.rgb
    this_run.font.highlight_color = tmpl_run.font.highlight_color
    this_run.font.strike = tmpl_run.font.strike

def put_html_text(document, docx_list):
    #We copy all the runs (with format) of the temporary docx with requirement text formated from html conversion into the main document
    for docx_html_id in docx_list:
        for tables in document.tables:
                for celda in tables._cells:
                    for p in celda.paragraphs:
                        breakfor = False
                        if docx_html_id in p.text:
                            for html_paragraph in docx_list[docx_html_id].paragraphs: 
                                if html_paragraph.style.name == 'List Bullet':
                                    new_paragraph = celda.add_paragraph(style='BulletList')
                                elif html_paragraph.style.name == 'List Number':
                                    new_paragraph = celda.add_paragraph(style='NumberList')
                                else: new_paragraph = celda.add_paragraph(style='No Spacing')
                                for run in html_paragraph.runs:
                                    if run.text != '':
                                        cloned_run = new_paragraph.add_run()
                                        clone_run_props(run, cloned_run)
                                        cloned_run.text = run.text
                            breakfor = True
                            #nexet we remove the placeholders
                            p_el = p._element
                            p_el.getparent().remove(p_el)
                            p._p = p._element = None
                    if breakfor:
                        break
                if breakfor:
                    break

def filter_specification(element):
    if element['specification'] == CURRENT_SPECIFICATION["id"]: 
        return True  
    return False

def filter_images(file):
    if file['mimetype'] == 'image/jpeg' or file['mimetype'] == 'image/png':
        return True  
    return False

def filter_not_images(file):
    if file['mimetype'] != 'image/jpeg' and file['mimetype'] != 'image/png':
        return True  
    return False

def has_images(requirement,images):
    for objectid in images:
        if images[objectid]['object_id'] == requirement:
            return True  
    return False

def get_requirement_images(requirement,images):
    requirement_images = {}
    for image in images:
        if images[image]['object_id'] == requirement:
            requirement_images[image] =  images[image]   
    return requirement_images

def get_requirement_applicability(requirement,all_project_component_types, all_project_applicability_conditions):
    all_requirement_applicabilities = ""

    for applicability in all_project_applicability_conditions:
        if all_project_applicability_conditions[applicability]['requirement'] == requirement:
            counter = 0
            for component_types in all_project_applicability_conditions[applicability]['component_types']:
                if len(all_project_applicability_conditions[applicability]['component_types']) == 1:
                    all_requirement_applicabilities+=all_project_component_types.get(component_types)['name']
                elif counter == 0:
                    all_requirement_applicabilities+=all_project_component_types.get(component_types)['name']
                else: all_requirement_applicabilities+="|"+all_project_component_types.get(component_types)['name']
                counter+= 1
            all_requirement_applicabilities+=" ; "
    if (len(all_requirement_applicabilities)) > 1:
        all_requirement_applicabilities = all_requirement_applicabilities[:-3]
    return all_requirement_applicabilities

def get_requirement_verification_methods(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['method'] != None:
                requirement_vms+=  vm['method']['name']
            requirement_vms+=" ; "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms

def get_requirement_verification_methods_newline(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['method'] != None:
                requirement_vms+=  vm['method']['name']
                requirement_vms+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms

def get_requirement_verification_methods_comments(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['component_vms'] != None:
                for component_vms in vm['component_vms']:
                    if component_vms['comment'] != None:
                        requirement_vms+=  component_vms['comment'][3:-4] #hardcode remove of <p> and </p> while clean_text is not available for this field
                        requirement_vms+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms    

def get_requirement_verification_methods_text(verification_methods):
    requirement_vms_text = ""
    if verification_methods == None:
        return requirement_vms_text 
    else:
        for vm in verification_methods:
            if vm['text'] != None:
                requirement_vms_text+=  vm['text'][3:-4] #hardcode remove of <p> and </p> while clean_text is not available for this field                
    if (len(verification_methods)) > 0:
        requirement_vms_text = requirement_vms_text[:-3]
    return requirement_vms_text    

def get_requirement_verification_methods_and_text(requirement,all_project_requirement_vms,all_project_requirement_ver_methods):
    requirement_vms = ""
    for req_vms_id in list(all_project_requirement_vms):
        req_vms = all_project_requirement_vms[req_vms_id]
        if req_vms['requirement'] == requirement:
            if req_vms['text'] != None:
                vm_method_and_text = req_vms['text'][:3] + all_project_requirement_ver_methods[req_vms['method']]['name'] +': ' + req_vms['text'][3:]
                requirement_vms+= vm_method_and_text
            all_project_requirement_vms.pop(req_vms_id) 
    return requirement_vms


def get_requirement_attachments_references(requirement,files):
    requirement_attachments = ""
    for file in files:
        if files[file]['object_id'] == requirement:
            if files[file]['reference'] != None:
                requirement_attachments+=  files[file]['name']+" - "+files[file]['reference'] #" - "+files[file]['version']   #for the future when we also upload the file versiopn and not controled by VS
            else :
                requirement_attachments+=  files[file]['name'] #+" - "+files[file]['reference'] #" - "+files[file]['version']   #for the future when we also upload the file versiopn and not controled by VS
            requirement_attachments+=" ; "
    if (len(requirement_attachments)) > 0:
        requirement_attachments = requirement_attachments[:-3]
    return requirement_attachments

def get_requirement_type(requirement,all_project_requirement_types):    
    for req_types in all_project_requirement_types:
        if all_project_requirement_types[req_types]['id'] == requirement['type']:
            return  all_project_requirement_types[req_types]['name']
    return ""

def get_requirement_owner(ownerId,all_deployment_users, all_deployment_user_groups):
    #It will get the User groups (removig Internal and External) + the user First and Last Names
    #needs to be fixed! 
    req_owner = "" 
    if ownerId == None:
        return req_owner
    else: req_owner_id = ownerId['id']

    req_owner = all_deployment_users[req_owner_id]['first_name']+ " " + all_deployment_users[req_owner_id]['last_name']
    req_owner_groups = ""

    if len(all_deployment_users[req_owner_id]['groups'])>0:
        for user_groups in all_deployment_users[req_owner_id]['groups']:
            if all_deployment_user_groups[user_groups]['name'] != "Internal" and all_deployment_user_groups[user_groups]['name'] != "External":
                req_owner_groups+=  all_deployment_user_groups[user_groups]['name']
                req_owner_groups+=" / "
        if (len(req_owner_groups)) > 0:
            req_owner_groups = req_owner_groups[:-3]
            req_owner = req_owner_groups + " -> " + req_owner
        
    return req_owner

def get_requirement_custom_field(requirement,all_project_custom_fields,all_project_custom_field_options,field_name):
    
    custom_field_value = ""
    req_custom_fields = requirement['custom_values']
    if len(req_custom_fields) > 0 :
        for custom_field in req_custom_fields:
            int_custom_field = int(custom_field)
            if all_project_custom_fields[int_custom_field]['name'] == field_name:
                if all_project_custom_fields[int_custom_field]['type'] == 0:
                    custom_field_value = req_custom_fields[custom_field]
                    custom_field_value = custom_field_value[3:-4] #to remove the <p> and </p>
                elif all_project_custom_fields[int_custom_field]['type'] == 1:
                    custom_field_value = all_project_custom_field_options[req_custom_fields[custom_field]]['name'] 
    return custom_field_value

def get_map(api: API, endpoint: Str = "/", name: Str = "id", name_map_func: Callable[[Str], Boolean] = None, filter_func: Callable[[Str], Boolean]= None):
    """
    Function that given an endpoint returns a dict with specific keys.
    If function is provided it generates the key. name_map_func must receive an object instance.
    Otherwise key will be the property of each object specified in name.
    """
    mapping = {}
    if not name_map_func:
        name_map_func = lambda x: x[name]
    for inst in api.get(endpoint):
        if filter_func and not filter_func(inst):
            # Filtered out
            continue

        key = name_map_func(inst)
        if not mapping.get(key):
            mapping[key] = inst
        else:
            warnings.warn(f"Warning ---> Key: {key} already has an object. Some data may be lost in mapping.")
    return mapping

def get_specification_requirements(all_project_requirements):
    spec_requirements = {}
    for requirement in all_project_requirements:
        if all_project_requirements[requirement]['specification'] == CURRENT_SPECIFICATION["id"]:
            spec_requirements[requirement] = all_project_requirements[requirement]
    return spec_requirements        

def create_file(api, file_name, file_data, project_id: int, obj=None):

    del api._session.headers['Content-Type']
    result = api._session.request("POST", api._url + "files/", files={
        'file': (file_name, file_data),
    }, data={
        'project': project_id,
 #       'folder': DEFAULT_VALUES['specs_folder'],
        'content_type': obj['contenttype'] if obj else 13,
        'object_id': obj['id'] if obj else project_id,
    })
    try:
        result.raise_for_status()
        new_file = result.json()
    except:
        new_file = None
    api._session.headers['Content-Type'] = "application/json"
    return new_file

def create_specification_document(api, document,all_project_requirements, specification_data, all_project_images, all_project_component_types, all_project_applicability_conditions,all_project_requirement_types,all_project_requirement_vms,all_project_requirement_ver_methods):
    
    print ("Going to create Specification document for -> "+ specification_data['name'])
    OUTPUT_FILE = specification_data['name']+'.docx' 
    CURRENT_SPECIFICATION["id"] = specification_data['id']

    template_data = []
   
    all_specification_requirements = get_specification_requirements(all_project_requirements)
    
    if len(all_specification_requirements) <1:
        print("No requirements for Specification -> "+ specification_data['name'])
        return 
        
        
    CURRENT_SPECIFICATION["name"] = specification_data['name']
    new_parser = HtmlToDocx()
    docx_list = {}
    #1st we will add requirements without section to the document
    no_section_requirements = get_requirements_without_section(all_specification_requirements)
    if len(no_section_requirements) >0:
        counter = 0
        for requirement in no_section_requirements:
            counter += 1
            reqdata = no_section_requirements[requirement]  

            req_vms = get_requirement_verification_methods(reqdata['verification_methods'])
            #req_vms_text = get_requirement_verification_methods_text(reqdata['verification_methods'])
            req_vms_text = get_requirement_verification_methods_and_text(requirement,all_project_requirement_vms,all_project_requirement_ver_methods)
           
            req_applicability = get_requirement_applicability(requirement,all_project_component_types, all_project_applicability_conditions)
            req_type = get_requirement_type(reqdata,all_project_requirement_types)

            requirement_with_images = has_images(requirement, all_project_images)     
            
            template_data.append({
                "specification_name" :  CURRENT_SPECIFICATION["name"] if counter == 1 else "",
                "section_name" :  "", 
                "req_id" :  reqdata['identifier'],
                "req_title" : reqdata['title'],
                "req_text" : reqdata['identifier']+"_docx",
                "req_type" : req_type,
                "req_rationale" :  reqdata['comment'],
                "req_ver_methods" :  req_vms,
                "req_ver_methods_text" :  req_vms_text,
                "req_ver_methods_and_text" :  reqdata['identifier']+"vms_docx",
                "req_applicability" :  req_applicability, 
                "images" :  "Images_Placeholder_"+str(requirement) if requirement_with_images == True else "No_Images"
            })
            
            docx_list[reqdata['identifier']+"_docx"] = new_parser.parse_html_string(reqdata['text'])   
            docx_list[reqdata['identifier']+"vms_docx"] = new_parser.parse_html_string(req_vms_text)    
        no_loose_requirements = 0
    else : no_loose_requirements = 1

    
    #we need to get the specification section in the correct order to be added to the document       
    specification_sections = get_specification_sections(api, DEFAULT_VALUES["project"])

    #for each section we will get its ordered requirements and add to the document
    for section in specification_sections:

        section_requirements = get_section_requirements(all_specification_requirements,section) 
        section_name = specification_sections[section]['name']
        counter = 0

        for requirement in section_requirements:
            counter += 1
            
            reqdata = section_requirements[requirement]
            
            req_vms = get_requirement_verification_methods(reqdata['verification_methods'])
            #req_vms_text = get_requirement_verification_methods_text(reqdata['verification_methods'])
            req_vms_text = get_requirement_verification_methods_and_text(requirement,all_project_requirement_vms,all_project_requirement_ver_methods)
           
            requirement_with_images = has_images(requirement, all_project_images)

            req_applicability = get_requirement_applicability(requirement,all_project_component_types, all_project_applicability_conditions)
            
            req_type = get_requirement_type(reqdata,all_project_requirement_types)

            template_data.append({
                "specification_name" :  CURRENT_SPECIFICATION["name"] if counter == 1 and no_loose_requirements == 1 else "",
                "section_name" :  section_name if counter == 1 else "", 
                "req_id" :  reqdata['identifier'],
                "req_title" : reqdata['title'],
                "req_text" : reqdata['identifier']+"_docx",
                "req_type" : req_type,
                "req_rationale" :  reqdata['comment'],
                "req_ver_methods" :  req_vms,
                "req_ver_methods_text" :  req_vms_text,
                "req_ver_methods_and_text" :  reqdata['identifier']+"vms_docx",
                "req_applicability" :  req_applicability, 
                "images" :  "Images_Placeholder_"+str(requirement) if requirement_with_images == True else "No_Images"
                })
            docx_list[reqdata['identifier']+"_docx"] = new_parser.parse_html_string(reqdata['text']) 
            docx_list[reqdata['identifier']+"vms_docx"] = new_parser.parse_html_string(req_vms_text)    
        no_loose_requirements = 0

    print ("Going to merge data into templates")
       
    document.merge_templates(template_data, separator='continuous_section')
    document.write(OUTPUT_FILE)

    document2 = Document(OUTPUT_FILE)
    remove_all_but_last_section(document2)
    remove_all_empty_headings(document2)
    print ("Going to put html formatting into document")
    put_html_text(document2, docx_list)
    print ("Going to put images into document")
    put_images(document2, all_project_images)   
    keep_tables_on_one_page(document2)

    document2.save(OUTPUT_FILE)
    create_file(api, OUTPUT_FILE, open(OUTPUT_FILE,'rb'), DEFAULT_VALUES["project"])    
    print ("Specification document created -> "+ specification_data['name'])
    return

def main(**kwargs):    
    #VALISPACE['domain'] = input("Enter your Valispace's deployment [https://....]:  ")
    #VALISPACE['username'] = input("Enter your username for Valispace's deployment:  ")
    #VALISPACE['password'] = getpass.getpass("Enter your password for Valispace's deployment:  ")    
    #DEFAULT_VALUES['project'] = input("Enter the project_ID to generate Specifications:  ")
    #DEFAULT_VALUES['files_directory'] = input("Enter your Template File Path (example c:/templates/): ")
    api = API(
            url = VALISPACE.get('domain'),
            username = VALISPACE.get('username'),
            password = VALISPACE.get('password'),
            warn_https = VALISPACE.get('warn_https', False),
        )
  

    all_project_specifications = get_map(api, f"requirements/specifications/?project="+str(DEFAULT_VALUES["project"]), "name")
    all_project_images = get_map(api, f"files/?project="+str(DEFAULT_VALUES["project"]), "id", None, filter_images)
    all_project_component_types = get_map(api, f"components/types/", "id", )
    all_project_applicability_conditions = get_map(api, f"requirements/applicability-conditions/?project="+str(DEFAULT_VALUES["project"]), "id")
    all_project_requirement_types = get_map(api, f"requirements/types/?project="+str(DEFAULT_VALUES["project"]), "id")
    all_project_requirement_ver_methods = get_map(api, f"requirements/verification-methods/?project="+str(DEFAULT_VALUES["project"]), "id")
    all_project_requirement_vms = get_map(api, f"requirements/requirement-vms/?project="+str(DEFAULT_VALUES["project"])+"&clean_html=text", "id")
    all_project_requirements = get_map(api, f"requirements/complete/?project="+str(DEFAULT_VALUES["project"])+"&clean_html=text&clean_text=comment", "id")
    

    #all_deployment_user_groups = get_map(api, f"group/", "id")
    #all_deployment_users = get_map(api, f"user/", "id")
    #all_project_files = get_map(api, f"files/?project="+str(DEFAULT_VALUES["project"]), "id", None, filter_not_images)
    #all_project_custom_fields = get_map(api, f"data/custom-fields/?project="+str(DEFAULT_VALUES["project"]), "id")
    #all_project_custom_field_options = get_map(api, f"data/custom-field-options/?project="+str(DEFAULT_VALUES["project"]), "id")
    import_file = api.get("files/189/")
    readfile = BytesIO(requests.get(import_file['download_url']).content)

    for specification in all_project_specifications:
        create_specification_document(api, MailMerge(readfile, remove_empty_tables=False, auto_update_fields_on_open="auto"),all_project_requirements, all_project_specifications[specification], all_project_images, all_project_component_types,all_project_applicability_conditions,all_project_requirement_types,all_project_requirement_vms,all_project_requirement_ver_methods)

    pass

if __name__=='__main__':
    main()
