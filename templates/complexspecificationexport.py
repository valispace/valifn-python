import docx
import operator
import warnings
import urllib.request
import ssl
import re


from ast import Str
from collections.abc import Callable
from xmlrpc.client import Boolean
import valispace
from PIL import Image
from docx import Document
from docx.shared import Inches
from mailmerge import MailMerge
from natsort import natsorted 
from io import BytesIO
import datetime
from docx.oxml.shared import qn
from htmldocx import HtmlToDocx
from docx.oxml.shared import OxmlElement


from .settings import Username, Password

#from .settings import Export_Project, Export_Create_Links, Export_Create_Index, Export_Template_Id, Export_Document_Name
#from .settings import Export_Custom_Column_One, Export_Custom_Column_Two, Export_Specification, Export_Requirement_States

Export_Specification = "Spacecraft_Spec"   
Export_Template_Id = 167
Export_Project = 24
Export_Document_Name = "Spec Export Document"
Export_Custom_Column_One = "none"
Export_Custom_Column_Two = "Protocol none"
Export_Requirement_States = "All"
Export_Create_Links = "True"
Export_Remove_Empty_Fields = "False"
Export_Create_Index = "False"


CURRENT_SPECIFICATION = {
    "id": 0,
    "name": "",
    "section": 0,
    "specs_folder": "",
}

DEFAULT_VALUES = {
    "domain": "https://demonstration.valispace.com/",
    "states_to_consider": []
}

def get_bookmark_par_element(document, bookmark_name):
    """
    Return the named bookmark parent paragraph element. If no matching
    bookmark is found, the result is '1'. If an error is encountered, '2'
    is returned.
    """
    doc_element = document.part.element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            par = bookmark.getparent()
            if not isinstance(par, docx.oxml.CT_P): 
                return 2
            else:
                return par
    return 1

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def create_reference_documents_index(document,all_specification_requirements, all_project_files):

    table = document.add_table(rows=1, cols=3, style="Style1")
    count = 0
    tableHasContent = False
    
    for requirement in all_specification_requirements:
        count += 1
        todeletefiles = []
        for requirement in all_specification_requirements:
            todeletefiles = []
            for file in all_project_files:
                if all_project_files[file]['object_id'] == requirement:                
                    count += 1
                    if all_project_files[file]['file_type'] == 1:
                        ref_file = all_project_files[file]
                    else: ref_file = all_project_files[all_project_files[file]['reference_file']] #it's external file ref

                    tableHasContent = True
                    todeletefiles.append(file)
                    if count == 1:
                        heading_row = table.rows[0].cells
                        heading_row[0].text = "FILE NAME"
                        heading_row[1].text = "FILE VERSION"
                        heading_row[2].text = "FILE REFERENCE"
                        data_row = table.add_row().cells
                        data_row[0].text = ref_file['name']
                        if ref_file['current_version'] != None:
                            data_row[1].text = str(ref_file['current_version'])
                        else: data_row[1].text = ""
                        if ref_file['reference'] != None:
                            data_row[2].text = ref_file['reference'][3:-4]  #we remove by "hand" the <p> and </p>
                        else: data_row[2].text = ""
                    else:
                        data_row = table.add_row().cells
                        data_row[0].text = ref_file['name']
                        if ref_file['current_version'] != None:
                            data_row[1].text = str(ref_file['current_version'])
                        else: data_row[1].text = ""
                        if ref_file['reference'] != None:
                            data_row[2].text = ref_file['reference'][3:-4]  #we remove by "hand" the <p> and </p>
                        else: data_row[2].text = ""
            for todeletefile in todeletefiles:
                del all_project_files[todeletefile]
    if tableHasContent :
        regexp = re.compile(CURRENT_SPECIFICATION["name"])
        for paragraph in document.paragraphs:
            if paragraph.text and regexp.search(paragraph.text):
                tbl, p = table._tbl, paragraph._p
                p.addnext(tbl) #moves index table from end of document to after the Specification Name
                tbl.addnext(p) #moves index table before the Specification Name
    else:
        table._element.getparent().remove(table._element)
    return

def put_url(document, req_url):
    #We copy all the runs (with format) of the temporary docx with requirement text formated from html conversion into the main document
    for docx_url_id in req_url:
        for tables in document.tables:
                for celda in tables._cells:
                    for p in celda.paragraphs:
                        breakfor = False
                        if docx_url_id in p.text:
                            hyperlink = add_hyperlink(p, req_url[docx_url_id], p.text[:-4], None, True)
                            breakfor = True
                            #nexet we remove the placeholders 
                            for run in reversed(list(p.runs)):
                                p._p.remove(run._r)
                    if breakfor:
                        break
                if breakfor:
                    break

def keep_tables_on_one_page(doc):
    allTables = doc.tables
    for activeTable in allTables:
        if activeTable.cell(0,0).paragraphs[0].text == 'REMOVE_TABLE':
            activeTable._element.getparent().remove(activeTable._element)
    
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
         ppr = tag.get_or_add_pPr()
         ppr.keepNext_val = True

    tags = doc.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag

def adjust_headings(document):
    pattern = "^\d+(?:\.\d+)*"
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Title" and paragraph.text =='': #removes empty headings
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
        elif paragraph.style.name == "Heading 1":            
            #Assuming Section are alway in format "number.number.number section title"
            clean_header_array = re.split(pattern,paragraph.text)
            if paragraph.text =='': #removes empty headings
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
                continue
            elif paragraph.text.count(".") == 0:
                paragraph.text = paragraph.text[2:]
            elif paragraph.text.count(".") == 1:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 2"
            elif paragraph.text.count(".") == 2:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 3"
            elif paragraph.text.count(".") == 3:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 4"
            elif paragraph.text.count(".") == 4:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 5"
            elif paragraph.text.count(".") == 5:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 6"
            elif paragraph.text.count(".") == 6:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 7"
            elif paragraph.text.count(".") == 7:
                if len(clean_header_array)>1:
                    paragraph.text = clean_header_array[1][1:]
                paragraph.style = "Heading 8"

def remove_all_but_last_section(document):
    sectPrs = document._element.xpath(".//w:pPr/w:sectPr")
    for sectPr in sectPrs:
        sectPr.getparent().remove(sectPr)
    return

def get_requirements_without_section(all_spec_requirements):
    unsorted_section_requirements = {}    
    no_section_requirements = {}
    for requirement in all_spec_requirements:
        if all_spec_requirements[requirement]['group'] == None:
            unsorted_section_requirements[requirement] =  all_spec_requirements[requirement]   

    sorted_requirements = natsorted(unsorted_section_requirements.items(), key=lambda x:(operator.getitem(x[1],'identifier').replace(". "," ")))
    
    for sorted_ in sorted_requirements:
        no_section_requirements[sorted_[0]] = sorted_[1] 

    return no_section_requirements
 
def get_section_requirements(all_specification_requirements,section):
    unsorted_section_requirements = {}
    section_requirements = {}
    for requirement in all_specification_requirements:
        if all_specification_requirements[requirement]['group'] == section:
            unsorted_section_requirements[requirement] =  all_specification_requirements[requirement]  
    
    sorted_requirements_by_section = natsorted(unsorted_section_requirements.items(), key=lambda x:(operator.getitem(x[1],'identifier').replace(". "," ")))
    
    for sorted_ in sorted_requirements_by_section:
        section_requirements[sorted_[0]] = sorted_[1] 

    return section_requirements

def get_specification_sections(api, project):
    specification_sections = {}
    all_specifications_groups = get_map(api, f"requirements/groups/?project={project}", "id", None, filter_specification)
    pre_sorted_section_groups = natsorted(all_specifications_groups.items(), key=lambda x:(operator.getitem(x[1],'name').replace(". "," ")))
    
    for pre_sorted in pre_sorted_section_groups:
        specification_sections[pre_sorted[0]] = pre_sorted[1] 

    return specification_sections

def put_images_and_remove_empty_fields(document, all_project_images):

    for paragraph in document.paragraphs:
        if "Images_Placeholder" in paragraph.text:
            #still needs testing/refined. Currently only working/tested for tables!
            req_images = get_requirement_images(requirement_id,all_project_images)
            for image in req_images:
                image_data = Image.open(BytesIO(urllib.request.urlopen(req_images[image]['download_url'], context=ssl.SSLContext()).read()))     
                run = paragraph.add_run()
                if image_data.info['dpi'][0] == 0:
                # Open the image using Pillow
                    with Image.open(image_data.fp) as im:
                     # Convert the image to RGB mode if it's not
                        if im.mode != 'RGB':
                            im = im.convert('RGB')
                        # Save the image to a BytesIO object, setting the DPI to a non-zero value
                        image_bytes = BytesIO()
                        im.save(image_bytes, format='JPEG', dpi=(96, 96))  # Set your desired DPI here
                        image_bytes.seek(0)
                        # Add the image to the docx file
                        run.add_picture(image_bytes, width=Inches(6))            
                elif image_data.width/image_data.info['dpi'][0] > celda.width.inches:
                    run.add_picture(image_data.fp, width=Inches(6))        
                else:
                    run.add_picture(image_data.fp)
                for run in reversed(list(paragraph.runs)[:-len(req_images)]):
                    paragraph._p.remove(run._r)
        elif "No_Images" in paragraph.text:
            p_el = paragraph._element
            p_el.getparent().remove(p_el)
            paragraph._p = paragraph._element = None
        elif "[TO_DELETE]" in paragraph.text:
            p_el = paragraph._element
            p_el.getparent().remove(p_el)
            paragraph._p = paragraph._element = None
        elif paragraph.text == "":
            p_el = paragraph._element
            if isinstance(p_el.getprevious() , docx.oxml.CT_P): 
                p_el.getparent().remove(p_el)
                paragraph._p = paragraph._element = None
                


    for tables in document.tables:
            for celda in tables._cells:
                hasImages = False
                requirement_id = 0
                for p in celda.paragraphs:
                    if "Images_Placeholder" in p.text:
                       requirement_id = int(p.text[19:])
                       hasImages = True                       
                       p_el = p._element
                       p_el.getparent().remove(p_el)
                       p._p = p._element = None
                    elif "No_Images" in p.text:
                        p_el = p._element
                        p_el.getparent().remove(p_el)
                        p._p = p._element = None
                    elif "[TO_DELETE]" in p.text:
                        c_el = celda._element
                        row_el = c_el.getparent()
                        row_el.getparent().remove(row_el) 
                if hasImages:
                    tables.autofit = False
                    new_image_p = celda.add_paragraph()
                    req_images = get_requirement_images(requirement_id,all_project_images)
                    for image in req_images:
                        image_data = Image.open(BytesIO(urllib.request.urlopen(req_images[image]['download_url'], context=ssl.SSLContext()).read())) 
                        run = new_image_p.add_run()
                        if image_data.info['dpi'][0] == 0:
                        # Open the image using Pillow
                            with Image.open(image_data.fp) as im:
                            # Convert the image to RGB mode if it's not
                                if im.mode != 'RGB':
                                    im = im.convert('RGB')
                                # Save the image to a BytesIO object, setting the DPI to a non-zero value
                                image_bytes = BytesIO()
                                im.save(image_bytes, format='JPEG', dpi=(96, 96))  # Set your desired DPI here
                                image_bytes.seek(0)
                                # Add the image to the docx file
                                run.add_picture(image_bytes, width=Inches(6))            
                        elif image_data.width/image_data.info['dpi'][0] > celda.width.inches:
                            run.add_picture(image_data.fp, width=Inches(6))        
                        else:
                            run.add_picture(image_data.fp)   
                    for run in reversed(list(new_image_p.runs)[:-len(req_images)]):
                        new_image_p._p.remove(run._r)
 
def clone_run_props(tmpl_run, this_run, original_run):
    this_run.bold = tmpl_run.bold
    this_run.italic = tmpl_run.italic
    this_run.underline = tmpl_run.underline       
    this_run.font.color.rgb = tmpl_run.font.color.rgb
    this_run.font.highlight_color = tmpl_run.font.highlight_color
    this_run.font.strike = tmpl_run.font.strike
    this_run.font.name = original_run.font.name 
    this_run.font.size = original_run.font.size 

def put_html_text(document, docx_list):
    #We copy all the runs (with format) of the temporary docx with requirement text formated from html conversion into the main document
    for docx_html_id in docx_list:
        for tables in document.tables:
                for celda in tables._cells:
                    for p in celda.paragraphs:
                        breakfor = False
                        if docx_html_id in p.text:
                            for html_paragraph in docx_list[docx_html_id].paragraphs: 
                                if html_paragraph.style.name == 'List Bullet':
                                    new_paragraph = celda.add_paragraph(style='BulletList')
                                elif html_paragraph.style.name == 'List Number':
                                    new_paragraph = celda.add_paragraph(style='NumberList')
                                else: new_paragraph = celda.add_paragraph(style=p.style)
                                for run in html_paragraph.runs:
                                    if run.text != '':
                                        cloned_run = new_paragraph.add_run()
                                        clone_run_props(run, cloned_run, p.runs[0])                                              
                                        cloned_run.text = run.text

                            breakfor = True
                            #nexet we remove the placeholders
                            p_el = p._element
                            p_el.getparent().remove(p_el)
                            p._p = p._element = None
                    if breakfor:
                        break
                if breakfor:
                    break
        for paragraph in document.paragraphs:
            if docx_html_id in paragraph.text:
                print(paragraph.text)
                for html_paragraph in docx_list[docx_html_id].paragraphs: 
                    print(html_paragraph)
                    if html_paragraph.style.name == 'List Bullet':
                        new_paragraph = paragraph.insert_paragraph_before(style='BulletList')
                    elif html_paragraph.style.name == 'List Number':
                        new_paragraph = paragraph.insert_paragraph_before(style='NumberList')
                    else: new_paragraph = paragraph.insert_paragraph_before(style=p.style)
                    for run in html_paragraph.runs:
                        if run.text != '':
                            cloned_run = new_paragraph.add_run()
                            clone_run_props(run, cloned_run, p.runs[0])
                            cloned_run.text = run.text
                p_el = paragraph._element
                p_el.getparent().remove(p_el)
                paragraph._p = paragraph._element = None        

def filter_specification(element):
    if element['specification'] == CURRENT_SPECIFICATION["id"]: 
        return True  
    return False

def filter_images(file):
    if file['mimetype'] == 'image/jpeg' or file['mimetype'] == 'image/png':
        return True  
    return False

def filter_not_images(file):
    if file['mimetype'] != 'image/jpeg' and file['mimetype'] != 'image/png':
        return True  
    return False

def has_images(requirement,images):
    for objectid in images:
        if images[objectid]['object_id'] == requirement:
            return True  
    return False

def get_requirement_images(requirement,images):
    requirement_images = {}
    for image in images:
        if images[image]['object_id'] == requirement:
            requirement_images[image] =  images[image]   
    return requirement_images

def get_requirement_applicability(requirement,all_project_component_types, all_project_applicability_conditions):
    all_requirement_applicabilities = ""

    for applicability in all_project_applicability_conditions:
        if all_project_applicability_conditions[applicability]['requirement'] == requirement:
            counter = 0
            for component_types in all_project_applicability_conditions[applicability]['component_types']:
                if len(all_project_applicability_conditions[applicability]['component_types']) == 1:
                    all_requirement_applicabilities+=all_project_component_types.get(component_types)['name']
                elif counter == 0:
                    all_requirement_applicabilities+=all_project_component_types.get(component_types)['name']
                else: all_requirement_applicabilities+="|"+all_project_component_types.get(component_types)['name']
                counter+= 1
            all_requirement_applicabilities+=" ; "
    if (len(all_requirement_applicabilities)) > 1:
        all_requirement_applicabilities = all_requirement_applicabilities[:-3]
    return all_requirement_applicabilities

def get_requirement_verification_methods(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['method'] != None:
                requirement_vms+=  vm['method']['name']
            requirement_vms+=" ; "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms

def get_requirement_verification_methods_newline(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['method'] != None:
                requirement_vms+=  vm['method']['name']
                requirement_vms+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms

def get_requirement_verification_methods_components_newline(verification_methods,all_project_components):
    requirement_vms_components = ""
    if verification_methods == None:
        return requirement_vms_components 
    else:
        for vm in verification_methods:
            if vm['component_vms'] != None:
                for component_vms in vm['component_vms']:
                    if component_vms['component'] != None:
                        requirement_vms_components+=  all_project_components[component_vms['component']]['name'] 
                        requirement_vms_components+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms_components = requirement_vms_components[:-3]
    return requirement_vms_components    

def get_requirement_verification_methods_status_newline(verification_methods):
    requirement_vms_status = ""
    if verification_methods == None:
        return requirement_vms_status 
    else:
        for vm in verification_methods:
            if vm['component_vms'] != None:
                for component_vms in vm['component_vms']:
                    if component_vms['status'] != None:
                        requirement_vms_status+=  component_vms['status'] 
                        requirement_vms_status+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms_status = requirement_vms_status[:-3]
    return requirement_vms_status 


def get_requirement_verification_methods_comments(verification_methods):
    requirement_vms = ""
    if verification_methods == None:
        return requirement_vms 
    else:
        for vm in verification_methods:
            if vm['component_vms'] != None:
                for component_vms in vm['component_vms']:
                    if component_vms['comment'] != None:
                        requirement_vms+=  component_vms['comment'][3:-4] #hardcode remove of <p> and </p> while clean_text is not available for this field
                        requirement_vms+=" \n\n "
    if (len(verification_methods)) > 0:
        requirement_vms = requirement_vms[:-3]
    return requirement_vms    

def get_requirement_verification_methods_text(verification_methods,all_project_requirement_vms):
    requirement_vms_text = ""
    if verification_methods == None:
        return requirement_vms_text 
    else:
        for vm in verification_methods:
            if vm['text'] != None:
                requirement_vms_text+=  all_project_requirement_vms[vm['id']]['text'][3:-4] #hardcode remove of <p> and </p> while clean_text is not available for this field                
    if (len(verification_methods)) > 0:
        requirement_vms_text = requirement_vms_text[:-3]
    return requirement_vms_text    

def get_requirement_verification_methods_and_text(requirement,all_project_requirement_vms,all_project_requirement_ver_methods):
    requirement_vms = ""
    for req_vms_id in list(all_project_requirement_vms):
        req_vms = all_project_requirement_vms[req_vms_id]
        if req_vms['requirement'] == requirement:
            if req_vms['text'] != None:
                vm_method_and_text = req_vms['text'][:3] + all_project_requirement_ver_methods[req_vms['method']]['name'] +': ' + req_vms['text'][3:]
                requirement_vms+= vm_method_and_text
            all_project_requirement_vms.pop(req_vms_id) 
    return requirement_vms

def get_requirement_attachments_references(requirement,files):
    requirement_attachments = ""
    for file in files:
        if files[file]['object_id'] == requirement:
            ref_file = files[file]
            if ref_file['file_type'] == 1:
                if ref_file['reference'] != None:
                    requirement_attachments+=  ref_file['name']+" - "+ref_file['reference'] #" - "+ref_file['version']   #for the future when we also upload the file versiopn and not controled by VS
                else :
                    requirement_attachments+=  ref_file['name'] #+" - "+ref_file['reference'] #" - "+ref_file['version']   #for the future when we also upload the file versiopn and not controled by VS
            elif ref_file['file_type'] == 3:
                original_ref_file = files[ref_file['reference_file']]
                if original_ref_file['reference'] != None:
                    requirement_attachments+=  original_ref_file['name']+" - "+original_ref_file['reference'] #" - "+original_ref_file['version']   #for the future when we also upload the file versiopn and not controled by VS
                else :
                    requirement_attachments+=  original_ref_file['name'] #+" - "+original_ref_file['reference'] #" - "+original_ref_file['version']   #for the future when we also upload the file versiopn and not controled by VS
            

            requirement_attachments+=" ; "
    if (len(requirement_attachments)) > 0:
        requirement_attachments = requirement_attachments[:-3]
    return requirement_attachments

def get_requirement_type(requirement,all_project_requirement_types):    
    for req_types in all_project_requirement_types:
        if all_project_requirement_types[req_types]['id'] == requirement['type']:
            return  all_project_requirement_types[req_types]['name']
    return ""

def get_requirement_tags(req_tags):
    tags = ""
    if req_tags == None:
        return tags 
    else:
        for tag in req_tags:
            tags+=  tag
            tags+=" ; "
    if (len(req_tags)) > 0:
        tags = tags[:-3]
    return tags

def get_requirement_owner(owner,all_deployment_users, all_deployment_user_groups):
    #It will get the User groups (removig Internal and External) + the user First and Last Names

    req_owner = "" 
    isSingleOwner = True
    if owner == None:
        return req_owner
    else: req_owner_id = owner['id']

    if owner['contenttype'] == 4 :
        isSingleOwner = False

    if isSingleOwner:    
        req_owner = all_deployment_users[req_owner_id]['first_name']+ " " + all_deployment_users[req_owner_id]['last_name']
        req_owner_groups = ""
        if len(all_deployment_users[req_owner_id]['groups'])>0:
            for user_groups in all_deployment_users[req_owner_id]['groups']:
                if all_deployment_user_groups[user_groups]['name'] != "Internal" and all_deployment_user_groups[user_groups]['name'] != "External":
                    req_owner_groups+=  all_deployment_user_groups[user_groups]['name']
                    req_owner_groups+=" / "
            if (len(req_owner_groups)) > 0:
                req_owner_groups = req_owner_groups[:-3]
                req_owner = req_owner_groups + " -> " + req_owner
    else:
        req_owner = all_deployment_user_groups[req_owner_id]['name']        
    return req_owner

def get_requirement_custom_field(requirement,all_project_custom_fields,all_project_custom_field_options,all_project_custom_field_rows,field_name):

    custom_field_value = ""
    custom_field = all_project_custom_fields.get(field_name)
    if custom_field == None :
        return custom_field_value
    
    custom_field_id = custom_field['id']
    req_custom_fields = requirement['custom_field_rows']
    if len(req_custom_fields) > 0 :
        for custom_field_row in req_custom_fields:
            if custom_field_row['custom_field'] == custom_field_id:
                if all_project_custom_fields[field_name]['type'] == 0:
                    custom_field_value = custom_field_row['value']
                    #custom_field_value = custom_field_value[3:-4] #to remove the <p> and </p>
                elif all_project_custom_fields[field_name]['type'] == 1:
                   if all_project_custom_fields[field_name]['multi']: 
                        for custom_field_option in custom_field_row['value']:
                            custom_field_value += all_project_custom_field_options[int(custom_field_option)]['name'] 
                            custom_field_value += " ; "
                        if (len(custom_field_value)) > 0:
                            custom_field_value = custom_field_value[:-3]
                   else: custom_field_value = all_project_custom_field_options[int(custom_field_row['value'])]['name'] 
    return custom_field_value

def get_requirement_parents(req_parents_ids,all_project_requirements):
    req_parents = ""
    for req_parent_id in req_parents_ids:
        req_parents += all_project_requirements[req_parent_id]['identifier'] + "\n"

    return req_parents

def get_requirement_children(req_children_ids,all_project_requirements):
    req_children = ""
    for req_children_id in req_children_ids:
        req_children += all_project_requirements[req_children_id]['identifier'] + "\n"

    return req_children

def filter_req_state(element):
    if element['state'] in DEFAULT_VALUES["states_to_consider"]: 
        return True  
    return False

def get_map(api, endpoint: Str = "/", name: Str = "id", name_map_func: Callable[[Str], Boolean] = None, filter_func: Callable[[Str], Boolean]= None):
    """
    Function that given an endpoint returns a dict with specific keys.
    If function is provided it generates the key. name_map_func must receive an object instance.
    Otherwise key will be the property of each object specified in name.
    """
    mapping = {}
    if not name_map_func:
        name_map_func = lambda x: x[name]
    for inst in api.get(endpoint):
        if filter_func and not filter_func(inst):
            # Filtered out
            continue

        key = name_map_func(inst)
        if not mapping.get(key):
            mapping[key] = inst
        else:
            warnings.warn(f"Warning ---> Key: {key} already has an object. Some data may be lost in mapping.")
    return mapping

def get_specification_requirements(all_project_requirements):
    spec_requirements = {}
    for requirement in all_project_requirements:
        if all_project_requirements[requirement]['specification'] == CURRENT_SPECIFICATION["id"]:
            spec_requirements[requirement] = all_project_requirements[requirement]
    return spec_requirements        

def helper_flag_line_for_deletion(value,Export_Remove_Empty_Fields):
    if value == "" and Export_Remove_Empty_Fields == "True":
        return "[TO_DELETE]"
    return value

def create_file(api, file_name, file_data, project_id: int, obj=None):

    del api._session.headers['Content-Type']
    result = api._session.request("POST", api._url + "files/", files={
        'file': (file_name, file_data),
    }, data={
        'project': project_id,
        'folder': CURRENT_SPECIFICATION['specs_folder'],
        'content_type': obj['contenttype'] if obj else 13,
        'object_id': obj['id'] if obj else project_id,
    })
    try:
        result.raise_for_status()
        new_file = result.json()
    except:
        new_file = None
    api._session.headers['Content-Type'] = "application/json"
    return new_file

def create_specification_document(api, Export_Remove_Empty_Fields, document,all_project_requirements, specification_data, all_project_images, all_project_component_types, all_project_applicability_conditions,all_project_requirement_types,all_project_requirement_vms,all_project_requirement_ver_methods,all_deployment_user_groups, all_deployment_users, all_project_files, all_project_custom_fields, all_project_custom_field_options,all_project_custom_field_rows,all_project_components):
    
    print ("Going to create Specification document for -> "+ specification_data['name'])
    OUTPUT_FILE = Export_Document_Name+'.docx' 
    CURRENT_SPECIFICATION["id"] = specification_data['id']

    template_data = []
   
    all_specification_requirements = get_specification_requirements(all_project_requirements)
    
    if len(all_specification_requirements) <1:
        print("No requirements for Specification -> "+ specification_data['name'])
        return 
        
        
    CURRENT_SPECIFICATION["name"] = specification_data['name']
    new_parser = HtmlToDocx()
    docx_list = {}
    req_url = {}
    #1st we will add requirements without section to the document
    no_section_requirements = get_requirements_without_section(all_specification_requirements)
    if len(no_section_requirements) >0:
        counter = 0
        for requirement in no_section_requirements:
            counter += 1
            reqdata = no_section_requirements[requirement]  
            requirement_with_images = has_images(requirement, all_project_images)  

            req_owner = get_requirement_owner(reqdata['owner'],all_deployment_users, all_deployment_user_groups)
            req_type = get_requirement_type(reqdata,all_project_requirement_types)
            req_children = get_requirement_children(reqdata['children'],all_project_requirements)
            req_parents = get_requirement_parents(reqdata['parents'],all_project_requirements)
            req_custom_column_one = get_requirement_custom_field(reqdata,all_project_custom_fields,all_project_custom_field_options,all_project_custom_field_rows,Export_Custom_Column_One)
            req_custom_column_two = get_requirement_custom_field(reqdata,all_project_custom_fields,all_project_custom_field_options,all_project_custom_field_rows,Export_Custom_Column_Two)
            
            req_vms = get_requirement_verification_methods_newline(reqdata['verification_methods'])
            
            req_vms_components = get_requirement_verification_methods_components_newline(reqdata['verification_methods'],all_project_components)
            req_vms_status = get_requirement_verification_methods_status_newline(reqdata['verification_methods'])
            req_tags = get_requirement_tags(reqdata['tags'])
            
            
            template_data.append({
                "specification_name" :  CURRENT_SPECIFICATION["name"] if counter == 1 else "",
                "section_name" :  "", 
                "req_id" :  reqdata['identifier']+"_url" if Export_Create_Links == "True" else reqdata['identifier'],
                "req_title" : reqdata['title'],
                "req_text" : reqdata['identifier']+"_docx",
                "req_type" : helper_flag_line_for_deletion(req_type ,Export_Remove_Empty_Fields),
                "req_rationale" :  helper_flag_line_for_deletion(reqdata['comment'] ,Export_Remove_Empty_Fields), 
                "req_owner" :  helper_flag_line_for_deletion(req_owner ,Export_Remove_Empty_Fields) ,
                "req_children" :  helper_flag_line_for_deletion(req_children ,Export_Remove_Empty_Fields),
                "req_parents" :  helper_flag_line_for_deletion(req_parents ,Export_Remove_Empty_Fields) ,
                "req_ver_methods" :  helper_flag_line_for_deletion(req_vms ,Export_Remove_Empty_Fields),
                "req_ver_status" :  helper_flag_line_for_deletion(req_vms_status ,Export_Remove_Empty_Fields),
                "req_ver_components" :  helper_flag_line_for_deletion(req_vms_components ,Export_Remove_Empty_Fields),
                "req_tags" :  helper_flag_line_for_deletion(req_tags ,Export_Remove_Empty_Fields),
                "images" :  "Images_Placeholder_"+str(requirement) if requirement_with_images == True else "No_Images"
            })
            
            docx_list[reqdata['identifier']+"_docx"] = new_parser.parse_html_string(reqdata['text'])   
            if Export_Create_Links == "True" : 
                req_url[reqdata['identifier']+"_url"]="https://"+DEFAULT_VALUES['domain']+"/project/"+str(Export_Project)+"/specifications/"+str(CURRENT_SPECIFICATION["id"])+"/requirements/"+str(reqdata['id'])
        no_loose_requirements = 0
    else : no_loose_requirements = 1

    
    #we need to get the specification section in the correct order to be added to the document       
    specification_sections = get_specification_sections(api, Export_Project)

    #for each section we will get its ordered requirements and add to the document
    for section in specification_sections:

        section_requirements = get_section_requirements(all_specification_requirements,section) 
        section_name = specification_sections[section]['name']
        counter = 0

        if len(section_requirements) == 0:
            template_data.append({
                "specification_name" : CURRENT_SPECIFICATION["name"] if counter == 1 and no_loose_requirements == 1 else "",
                "section_name" :  section_name,
                "req_id": "REMOVE_TABLE" 
            })

        for requirement in section_requirements:
            counter += 1
            
            reqdata = section_requirements[requirement]  
            requirement_with_images = has_images(requirement, all_project_images)  

            req_owner = get_requirement_owner(reqdata['owner'],all_deployment_users, all_deployment_user_groups)
            req_type = get_requirement_type(reqdata,all_project_requirement_types)
            req_children = get_requirement_children(reqdata['children'],all_project_requirements)
            req_parents = get_requirement_parents(reqdata['parents'],all_project_requirements)
            req_custom_column_one = get_requirement_custom_field(reqdata,all_project_custom_fields,all_project_custom_field_options,all_project_custom_field_rows,Export_Custom_Column_One)
            req_custom_column_two = get_requirement_custom_field(reqdata,all_project_custom_fields,all_project_custom_field_options,all_project_custom_field_rows,Export_Custom_Column_Two)
            
            req_vms = get_requirement_verification_methods_newline(reqdata['verification_methods'])
            
            req_vms_components = get_requirement_verification_methods_components_newline(reqdata['verification_methods'],all_project_components)
            req_vms_status = get_requirement_verification_methods_status_newline(reqdata['verification_methods'])
            req_tags = get_requirement_tags(reqdata['tags'])
            
            
            template_data.append({
                "specification_name" : CURRENT_SPECIFICATION["name"] if counter == 1 and no_loose_requirements == 1 else "",
                "section_name" :  section_name if counter == 1 else "", 
                "req_id" :  reqdata['identifier']+"_url" if Export_Create_Links == "True" else reqdata['identifier'],
                "req_title" : reqdata['title'],
                "req_text" : reqdata['identifier']+"_docx",
                "req_type" : helper_flag_line_for_deletion(req_type ,Export_Remove_Empty_Fields),
                "req_rationale" :  helper_flag_line_for_deletion(reqdata['comment'] ,Export_Remove_Empty_Fields), 
                "req_owner" :  helper_flag_line_for_deletion(req_owner ,Export_Remove_Empty_Fields) ,
                "req_children" :  helper_flag_line_for_deletion(req_children ,Export_Remove_Empty_Fields),
                "req_parents" :  helper_flag_line_for_deletion(req_parents ,Export_Remove_Empty_Fields) ,
                "req_ver_methods" :  helper_flag_line_for_deletion(req_vms ,Export_Remove_Empty_Fields),
                "req_ver_status" :  helper_flag_line_for_deletion(req_vms_status ,Export_Remove_Empty_Fields),
                "req_ver_components" :  helper_flag_line_for_deletion(req_vms_components ,Export_Remove_Empty_Fields),
                "req_tags" :  helper_flag_line_for_deletion(req_tags ,Export_Remove_Empty_Fields),
                "images" :  "Images_Placeholder_"+str(requirement) if requirement_with_images == True else "No_Images"
            })
            
            docx_list[reqdata['identifier']+"_docx"] = new_parser.parse_html_string(reqdata['text'])     
            if Export_Create_Links == "True" : 
                req_url[reqdata['identifier']+"_url"]="https://"+DEFAULT_VALUES['domain']+"/project/"+str(Export_Project)+"/specifications/"+str(CURRENT_SPECIFICATION["id"])+"/requirements/"+str(reqdata['id'])
        no_loose_requirements = 0

    print ("Going to merge data into templates")
       
    document.merge_templates(template_data, separator='continuous_section')
    document.write(OUTPUT_FILE)

    document2 = Document(OUTPUT_FILE) 
    remove_all_but_last_section(document2)
    print ("Going to put html formatting into document")
    put_html_text(document2, docx_list)    
    adjust_headings(document2)
    
    if Export_Create_Links == "True":
        print ("Going to put Requirement Urls")
        put_url(document2, req_url)  
        
    if Export_Create_Index == "True":
        create_reference_documents_index(document2,all_specification_requirements, all_project_files)

    print ("Going to put Images into document")
    put_images_and_remove_empty_fields(document2, all_project_images)   
    keep_tables_on_one_page(document2)        

    document2.save(OUTPUT_FILE)    
    
    create_file(api, OUTPUT_FILE, open(OUTPUT_FILE,'rb'), Export_Project)    
    print ("Specification document created -> "+ specification_data['name'])
    return

def main(**kwargs):    
    
    api = valispace.API(url=DEFAULT_VALUES["domain"],username=Username,password=Password,warn_https=False)

    all_project_requirement_states = get_map(api, f"vstates/?project="+str(Export_Project), "name")
    req_states_to_consider = []

    for req_state in all_project_requirement_states:
        if all_project_requirement_states[req_state]["name"] in Export_Requirement_States:
            req_states_to_consider.append(all_project_requirement_states[req_state]['id'])

    DEFAULT_VALUES['states_to_consider'] = req_states_to_consider

    all_project_specifications = get_map(api, f"requirements/specifications/?project="+str(Export_Project), "name")
    all_project_images = get_map(api, f"files/?project="+str(Export_Project), "id", None, filter_images)
    all_project_component_types = get_map(api, f"components/types/", "id")
    all_project_components = get_map(api, f"components/?project="+str(Export_Project), "id")
    all_project_applicability_conditions = get_map(api, f"requirements/applicability-conditions/?project="+str(Export_Project), "id")
    all_project_requirement_types = get_map(api, f"requirements/types/?project="+str(Export_Project), "id")
    all_project_requirement_ver_methods = get_map(api, f"requirements/verification-methods/?project="+str(Export_Project), "id")
    all_project_requirement_vms = get_map(api, f"requirements/requirement-vms/?project="+str(Export_Project)+"&clean_html=text", "id")
    all_deployment_user_groups = get_map(api, f"group/", "id")
    all_deployment_users = get_map(api, f"user/", "id")
    all_project_files = get_map(api, f"files/?project="+str(Export_Project), "id")
    all_project_custom_fields = get_map(api, f"data/custom-fields/?project="+str(Export_Project), "name")
    all_project_custom_field_options = get_map(api, f"data/custom-field-options/?project="+str(Export_Project), "id")
    all_project_custom_field_rows = get_map(api, f"data/custom-field-row/?project="+str(Export_Project), "id")
    if Export_Requirement_States != "All":
        all_project_requirements = get_map(api, f"requirements/complete/?project="+str(Export_Project)+"&clean_html=text&clean_text=comment", "id", None,filter_req_state)
    else:
        all_project_requirements = get_map(api, f"requirements/complete/?project="+str(Export_Project)+"&clean_html=text&clean_text=comment", "id")
    
    
    import_file = api.get(f"files/{Export_Template_Id}/")
    
    readfile = BytesIO(urllib.request.urlopen(import_file['download_url'], context=ssl.SSLContext()).read())

    specifications_folder = api.post('files/folders/', data={"project": Export_Project,"name": "ExportedSpecifications_"+datetime.datetime.now().strftime("%d-%b-%Y_%H:%M:%S")})
    CURRENT_SPECIFICATION['specs_folder'] = specifications_folder['id']

    
    if Export_Specification != "All":
        create_specification_document(api, Export_Remove_Empty_Fields,MailMerge(readfile, remove_empty_tables=False, auto_update_fields_on_open="auto"),all_project_requirements, all_project_specifications[Export_Specification], all_project_images, all_project_component_types,all_project_applicability_conditions,all_project_requirement_types,all_project_requirement_vms,all_project_requirement_ver_methods,all_deployment_user_groups, all_deployment_users, all_project_files, all_project_custom_fields, all_project_custom_field_options,all_project_custom_field_rows,all_project_components)
    else:
        for specification in all_project_specifications:
            create_specification_document(api,Export_Remove_Empty_Fields, MailMerge(readfile, remove_empty_tables=False, auto_update_fields_on_open="auto"),all_project_requirements, all_project_specifications[specification], all_project_images, all_project_component_types,all_project_applicability_conditions,all_project_requirement_types,all_project_requirement_vms,all_project_requirement_ver_methods,all_deployment_user_groups, all_deployment_users, all_project_files, all_project_custom_fields, all_project_custom_field_options,all_project_custom_field_rows,all_project_components)
   
    pass

if __name__=='__main__':
    main()
